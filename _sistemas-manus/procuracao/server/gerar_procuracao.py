#!/usr/bin/env python3
"""
gerar_procuracao.py — Gera o documento Word da procuração com os dados fornecidos.

Uso:
    python3 gerar_procuracao.py '<json_dados>' <caminho_modelo> <caminho_saida>
"""

import sys
import json
import os
import copy
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

MESES = [
    "", "janeiro", "fevereiro", "março", "abril", "maio", "junho",
    "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"
]


def data_por_extenso():
    hoje = datetime.now()
    return f"{hoje.day} de {MESES[hoje.month]} de {hoje.year}"


def substituir_no_paragrafo(paragraph, marcador, valor, forcar_bold=False):
    """
    Substitui `marcador` por `valor` num parágrafo, mesmo que o marcador
    esteja fragmentado entre múltiplos runs. Preserva a formatação do
    primeiro run que contém o início do marcador.
    """
    # Tenta substituição simples (marcador em um único run)
    for run in paragraph.runs:
        if marcador in run.text:
            run.text = run.text.replace(marcador, valor)
            if forcar_bold:
                run.bold = True
            return

    # Marcador fragmentado: reconstrói o texto completo do parágrafo
    texto_completo = "".join(r.text for r in paragraph.runs)
    if marcador not in texto_completo:
        return

    # Substitui no texto completo
    texto_novo = texto_completo.replace(marcador, valor)

    # Coloca tudo no primeiro run, limpa os demais
    if paragraph.runs:
        primeiro_run = paragraph.runs[0]
        # Preserva a formatação do run que continha o início do marcador
        pos = texto_completo.find(marcador[0])
        acum = 0
        run_ref = paragraph.runs[0]
        for r in paragraph.runs:
            if acum + len(r.text) > pos:
                run_ref = r
                break
            acum += len(r.text)

        # Copia formatação do run de referência para o primeiro
        primeiro_run.text = texto_novo
        if forcar_bold:
            primeiro_run.bold = True
        elif run_ref.bold is not None:
            primeiro_run.bold = run_ref.bold

        for r in paragraph.runs[1:]:
            r.text = ""


def gerar(dados: dict, modelo_path: str, saida_path: str):
    nome = dados["nome"].strip().upper()
    genero = dados["genero"].strip()
    estado_civil = dados["estadoCivil"].strip()
    profissao = dados["profissao"].strip()
    cpf = dados["cpf"].strip()
    naturalidade = dados["naturalidade"].strip()
    filiacao = dados["filiacao"].strip()
    rua = dados["rua"].strip()
    quadra = dados.get("quadra", "").strip()
    lote = dados.get("lote", "").strip()
    cidade = dados["cidade"].strip()
    cep = dados["cep"].strip()
    data = data_por_extenso()

    # Monta endereço
    endereco = rua
    if quadra:
        endereco += f", Quadra {quadra}"
    if lote:
        endereco += f", Lote {lote}"
    if cidade:
        endereco += f", {cidade}"
    if cep:
        endereco += f", CEP {cep}"

    # Lista de substituições: (marcador, valor, forcar_bold)
    substituicoes = [
        ("(NOME)",                     nome,         True),
        ("brasileiro (a)",             genero,       False),
        ("(Estado Civil)",             estado_civil, False),
        ("(profissão)",                profissao,    False),
        ("(xxx.xxx.xxx-xx)",           cpf,          False),
        ("(naturalidade = cidade-UF)", naturalidade, False),
        ("(filiação)",                 filiacao,     False),
        ("(endereço)",                 endereco,     False),
        ("(data atual)",               data,         False),
        ("(Assinatura do outorgante)", nome,         True),
    ]

    doc = Document(modelo_path)

    for paragraph in doc.paragraphs:
        for marcador, valor, bold in substituicoes:
            substituir_no_paragrafo(paragraph, marcador, valor, bold)

    # Verifica tabelas (se houver)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for marcador, valor, bold in substituicoes:
                        substituir_no_paragrafo(paragraph, marcador, valor, bold)

    doc.save(saida_path)
    print(f"OK:{saida_path}")


if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Uso: python3 gerar_procuracao.py '<json>' <modelo> <saida>", file=sys.stderr)
        sys.exit(1)

    dados_json = sys.argv[1]
    modelo_path = sys.argv[2]
    saida_path = sys.argv[3]

    try:
        dados = json.loads(dados_json)
    except json.JSONDecodeError as e:
        print(f"ERRO_JSON:{e}", file=sys.stderr)
        sys.exit(1)

    if not os.path.exists(modelo_path):
        print(f"ERRO_MODELO:Arquivo não encontrado: {modelo_path}", file=sys.stderr)
        sys.exit(1)

    try:
        gerar(dados, modelo_path, saida_path)
    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"ERRO:{e}", file=sys.stderr)
        sys.exit(1)
